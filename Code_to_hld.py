from flask import Flask, request, send_file, render_template_string,jsonify,Blueprint
from datetime import datetime
import json
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import logging
import base64
from PIL import Image
import urllib.parse
import urllib.request
import os
import time
import vertexai
import sys
import sqlite3
import uuid
from vertexai.preview.generative_models import GenerativeModel, Part
from vertexai.language_models import TextGenerationModel
import vertexai.preview.generative_models as generative_models
from flask_cors import CORS
import psycopg2
from psycopg2 import sql
from psycopg2.extensions import ISOLATION_LEVEL_AUTOCOMMIT
 
app = Flask(__name__)
CORS(app)
 
with open('postgrace_credential.json', 'r') as config_file:
    config = json.load(config_file)
    dbname = config.get('dbname')
    user = config.get('user')
    password = config.get('password')
    host = config.get('host')
    port = config.get('port')
 
 
 
target_path = "./backend_1/Logs"
 
# Initialize logger
logs_dir = os.path.join(target_path, 'Code_to_HLD_logs')
if not os.path.exists(logs_dir):
    os.makedirs(logs_dir)
    print(f"Code_to_HLD_logs directory created successfully at '{logs_dir}'.")
   
application_name = "Code_to_HLD_logs"
log_file_path = os.path.join(logs_dir, f"{datetime.now().strftime('%Y-%m-%d')}_{application_name}")
if not os.path.exists(log_file_path):
    os.makedirs(log_file_path)
    print(f"Code_to_HLD_logs directory created successfully at '{log_file_path}'.")
 
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
 
# Create handlers for each log level
file_handlers = {}
levels = ['error', 'info', 'debug', 'critical']
for level in levels:
    file_handler = logging.FileHandler(os.path.join(log_file_path, f"{level}.log"))
    file_handler.setLevel(logging.getLevelName(level.upper()))
    file_handler.setFormatter(formatter)
    file_handlers[level] = file_handler
 
    logger.addHandler(file_handler)
 
 
with open('config.json', 'r') as f:
    config_data = f.read()  # Read the JSON content as a string
    config = json.loads(config_data) # Converts string to dictionary
    project_id = config.get('project_id')
 
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "config.json"
# project_id = project_id
print(project_id)
vertexai.init(project=project_id, location="asia-southeast1")
 
 
def get_tsg_functional_response(buisness_use_case):
    config = {
        "max_output_tokens": 2048,
        "temperature": 0.9,
        "top_p": 1
    }
    model = GenerativeModel("gemini-pro")
    chat = model.start_chat()
    ui_response = chat.send_message(f"""\"{buisness_use_case}"
what all things are mandatory to be part of buisness story from this use case in user Interface test cases\n(give answer in 5 numbered points in new line)\n\nAnswer:""", generation_config=config, safety_settings={
    generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
})
    print(ui_response.text)
    usability_response = chat.send_message(f"""\"{buisness_use_case}"
what all things are mandatory to be part of buisness story from this use case in Usability test cases\n(give answer in 5 numbered points in new line)\n\nAnswer:""", generation_config=config, safety_settings={
    generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
})
    print(usability_response.text)
    functionality_response = chat.send_message(f"""\"{buisness_use_case}"
what all things are mandatory to be part of buisness story from this use case in Functionality test cases\n(give answer in 5 numbered points in new line)\n\nAnswer:""", generation_config=config, safety_settings={
    generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
})
    print(functionality_response.text)
    integration_response = chat.send_message(f"""\"{buisness_use_case}"
what all things are mandatory to be part of buisness story from this use case in Integration test cases\n(give answer in 5 numbered points in new line)\n\nAnswer:""", generation_config=config, safety_settings={
    generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
})
    print(integration_response.text)
    try:
        ui_response = ui_response.text
        usability_response = usability_response.text
        functionality_response = functionality_response.text
        integration_response = integration_response.text
        return f"### User interface (UI) test cases:\n{ui_response}\n\n### Usability test cases:\n{usability_response}\n\n### Functionality test cases:\n{functionality_response}\n\n### Integration test cases:\n{integration_response}\n\n"
    except KeyError as e:
        print("KeyError:", e)
        return "An error occurred while processing the response from OpenAI."
 
def get_tsg_nonfunctional_response(buisness_use_case):
    config = {
        "max_output_tokens": 2048,
        "temperature": 0.9,
        "top_p": 1
    }
    model = GenerativeModel("gemini-pro")
    chat = model.start_chat()
    security_response = chat.send_message(f"""\"{buisness_use_case}"
what all things are mandatory to be part of buisness story from this use case in Security test cases\n(give answer in 5 numbered points in new line)\n\nAnswer:""", generation_config=config, safety_settings={
    generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
})
    print(security_response.text)
    performance_response = chat.send_message(f"""\"{buisness_use_case}"
what all things are mandatory to be part of buisness story from this use case in Performance test cases\n(give answer in 5 numbered points in new line)\n\nAnswer:""", generation_config=config, safety_settings={
    generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
})
    print(performance_response.text)
    try:
        security_response = security_response.text
        performance_response = performance_response.text
        return f"### Security test cases:\n{security_response}\n\n### Performance test cases:\n{performance_response}\n\n"
    except KeyError as e:
        print("KeyError:", e)
        return "An error occurred while processing the response from GCP."
 
 
 
def web_sequence_diagram_code(code):
    config = {
        "max_output_tokens": 2048,
        "temperature": 0.9,
        "top_p": 1
    }
    model = GenerativeModel("gemini-pro")
    vertexai.init(project=project_id, location="asia-southeast1")
    chat = model.start_chat()
 
    code = chat.send_message(f"""\"{code}"
    generate code for web sequence diagram of above code .""", generation_config=config, safety_settings={
        generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
    })
 
    return code.text
 
 
 
 
def create_word_document(markdown_text, image_path=None):
    doc = Document()
    markdown_text = markdown_text.replace('**', '')  # Remove all occurrences of '**'
    markdown_lines = markdown_text.splitlines()
 
    for line in markdown_lines:
        if line.startswith("# "):
            # Heading 1
            heading_text = line.lstrip("# ")
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(18)  # Larger font size for Heading 1
        elif line.startswith("## "):
            # Heading 2
            heading_text = line.lstrip("## ")
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(16)  # Smaller font size for Heading 2
        elif line.startswith("### "):
            # Heading 3
            heading_text = line.lstrip("### ")
            p = doc.add_paragraph(heading_text)
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)  # Default font size for Heading 3
        elif line.startswith("#### "):
            # Heading 4
            heading_text = line.lstrip("### ")
            p = doc.add_paragraph(heading_text)
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(12)  # Default font size for Heading 4
        else:
            # Normal text
            doc.add_paragraph(line)
 
    if image_path:
        doc.add_paragraph()  # Add an empty paragraph
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(image_path, width=Inches(4), height=Inches(2))
 
    return doc
 
def getSequenceDiagram(text, outputFile, style='default'):
    request = {}
    request["message"] = text
    request["style"] = style
    request["apiVersion"] = "1"
 
    url = urllib.parse.urlencode(request).encode('utf-8')  # Encode the data and convert to bytes
 
    with urllib.request.urlopen("https://www.websequencediagrams.com/", url) as f:
        line = f.read().decode('utf-8')
 
    expr = re.compile("(\?(img|pdf|png|svg)=[a-zA-Z0-9]+)")
    m = expr.search(line)
 
    if m is None:
        print("Invalid response from server.")
        return False
 
    urllib.request.urlretrieve("https://www.websequencediagrams.com/" + m.group(0), outputFile)
    return True
 
def generate_data_from_code(code):
    input_form_prompt = code
 
    config = {
        "max_output_tokens": 2048,
        "temperature": 0.9,
        "top_p": 1
    }
    model = GenerativeModel("gemini-pro")
    chat = model.start_chat()
    response = chat.send_message(f"""\"{input_form_prompt}"
Here I will give you code as input to generate a well-structured and high level design document with all the headings with bold letters that I have defined from \n\n to \n\n8.. :\n\n- Heading: about the code in short form or 5 words \n\n1- Scope: Provide a concise five points explanation of the code's scope, outlining its intended purpose and functionality.\n\n2- Key Features: Identify the key features of the code and provide in-depth explanations for each feature, with at least three lines dedicated to each.\n\n3- Design Consideration: Elaborate on the code's design, discussing its features and functions by referencing the code itself. Provide a step-by-step explanation of how the entire process flows.\n\n4- Data Structure Explanation:\n##Extracted Information: Extract essential information such as the database name, table name, data types, and the rationale behind their selection.\n##ER Sequence Flow:Perform a sequence of actions for a data management scenario. The scenario may or may not involve a database. Follow the steps below:\n\n### Initial Setup\n   - [Specify any initial setup or fixture loading if applicable]\n\n### Entity Creation\n   - Create base records for one or more entities (e.g., Schools, Teachers, EmailLogins, AuditLogs).\n   - Establish relationships between entities if necessary.\n\n### Confirmation and Action\n   - Ask for confirmation or user input to trigger a specific action.\n   - Perform an action, such as resetting authentication or updating records.\n\n### Assertion\n   - Assert the state of the system after the action, checking record counts or specific conditions.\n\n### Exception Handling\n   - Include a scenario where an exception is expected.\n   - Specify the nature of the exception and how the system should respond.\n\n5- ER Sequence Diagram: Generate a workflow containing different nodes and their flow descriptions for the code provided. The generated sequence should strictly follow the pattern of 'One Node -> Another Node: Description' and include relevant test functions and descriptions that can be derived from the code. The goal is to identify and describe various nodes and workflow between those nodes. The output should be in the format of newline-separated text, and it should be structured and descriptive. Ensure that the generated nodes and descriptions are relevant to the code and capture the essence of the workflow of code. Starting with triple backticks (```) and ending with triple backticks (```)\ndo it step by step with the format in mind. Don't write if not in format!\nFormat- ```\nA Node ->Another Node: Data Flow Description\nA Node ->Another Node: Data Flow Description\n```\n\n6- UI Requirements Explanation: Provide a detailed explanation of various user interface (UI) requirements that have been considered, supporting your explanations with examples of functions.\n\n7- Security Requirements Explanation: \n##Security Requirements:- Identify and explain the security requirements that have been taken into account, and how they impact the overall flow. Reference the functions in the code as needed.\n##Security Design Principles:- Mention five Security Design principle by default.\n\n8- Integration Explanation:\n##Dependent Systems: List the dependent systems that the code interfaces with.\n##Connections and Data Flow: Describe the connections and data flow between the code and these dependent systems.
""", generation_config=config, safety_settings={
    generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
          generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
})
 
    response_text = response.text
    tsg_functional_response = get_tsg_functional_response(code)
    tsg_nonfunctional_response = get_tsg_nonfunctional_response(code)
 
    # Generating a sequence diagram from the code
    sequence_diagram_code_result = web_sequence_diagram_code(code)
    sequence_diagram_output_file = "web_sequence_diagram.png"
    getSequenceDiagram(sequence_diagram_code_result, sequence_diagram_output_file)
 
    # Creating a Word document with response text, functional response, and sequence diagram
    doc = create_word_document(f"{response_text}\n{tsg_functional_response}\n{tsg_nonfunctional_response}", sequence_diagram_output_file)
    doc_path = "Code_to_HLD.docx"
    doc.save(doc_path)
   
    with open(doc_path, "rb") as f:
        doc_data = f.read()
 
    b64 = base64.b64encode(doc_data).decode()
 
    # Construct the data URI for downloading the document
    href = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}"
 
    print("File downloaded successfully.")
    return href
 
@app.route('/codeToHLD', methods=['POST'])
def code_HLD():
    data = request.get_json()
    code = data.get('code')
    if code == "":
        logger.error("Write the code first. Without code, you cannot do analysis.")
        return "Write the code first. Without code, you cannot do analysis."
    new_uuid = uuid.uuid4()
    start_time = time.time()
    href = generate_data_from_code(code)
    end_time = time.time()
    total_time = round(end_time-start_time)
    try:
        conn = psycopg2.connect(
            dbname=dbname,
            user=user,
            password=password,
            host=host,
            port = port
        )
        cursor = conn.cursor()
 
        cursor.execute('''CREATE TABLE IF NOT EXISTS codehld (
                                key TEXT PRIMARY KEY,
                                month TEXT,
                                atdate DATE,
                                attime TIMESTAMP,
                                downloadlink TEXT,
                                total_time INTEGER
                    )''')
        conn.commit()
        atTime = datetime.now()
        month = datetime.now().strftime("%B")
        atDate = datetime.now().date()
        cursor.execute('''INSERT INTO codehld
                        (key, month, atdate, attime, downloadlink, total_time)
                        VALUES (%s, %s, %s, %s, %s, %s)''',
                        (str(new_uuid), str(month), str(atDate), str(atTime), href, total_time))
        conn.commit()
        print("Data inserted successfully into the table")
 
    except psycopg2.Error as e:
        print("An error occurred:", e)
    finally:
        if conn is not None:
            conn.close()  
    return jsonify({"href":href})
@app.route('/getEntireTableCodeHLD', methods=['GET'])
def get_entire_tabletranscribe():
    try:
        conn = psycopg2.connect(
            dbname=dbname,
            user=user,
            password=password,
            host=host,
            port = port
        )
 
        table_name = "codehld"
        cursor.execute("SELECT * FROM codehld ORDER BY attime DESC, atdate DESC")
        rows = cursor.fetchall()
 
        cursor.execute("SELECT COUNT(*) FROM codehld")
        count = cursor.fetchone()[0]
 
        cursor.execute("SELECT column_name FROM information_schema.columns WHERE table_name = %s", (table_name,))
        columns = cursor.fetchall()
        column_names = [col[0] for col in columns]
 
        print(count)
        conn.close()
 
        data = []
        for row in rows:
            row_dict = {}
            for idx, column_name in enumerate(column_names):
                row_dict[column_name] = row[idx]
            data.append(row_dict)
 
        print("Data fetched successfully")
        return jsonify({"data": data})
 
    except psycopg2.Error as e:
        print("An error occurred:", e)
        return jsonify({'error': 'An error occurred while fetching data from the database'}), 500
 
if __name__ == "__main__":
    app.run(host='0.0.0.0', debug=True, port=6500)